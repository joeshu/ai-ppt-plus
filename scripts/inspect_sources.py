#!/usr/bin/env python3
"""Inventory presentation sources without mutating them.

Usage: inspect_sources.py INPUT... --output report.json
Output is JSON; exit 0 readable, 2 missing/unreadable, 3 runtime failure.
The command is idempotent and only replaces its report. Optional dependencies:
PyMuPDF, python-docx, openpyxl and Pillow. Example/test: python inspect_sources.py . -o sources.json
"""
import argparse, hashlib, json, mimetypes, zipfile
from datetime import datetime, timezone
from pathlib import Path

def sha256(path):
    h=hashlib.sha256()
    with path.open('rb') as f:
        for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
    return h.hexdigest()

def inspect(path):
    r={'path':str(path.resolve()),'name':path.name,'size':path.stat().st_size,'empty':path.stat().st_size==0,'sha256':sha256(path),'kind':path.suffix.lower().lstrip('.') or 'unknown','mime':mimetypes.guess_type(path.name)[0],'readable':True,'details':{},'needs_ocr':False,'sensitive_signals':[],'errors':[]}
    try:
        ext=path.suffix.lower()
        if ext in {'.pptx','.docx','.xlsx'}:
            r['details']['valid_zip']=zipfile.is_zipfile(path)
            if not r['details']['valid_zip']: raise ValueError('invalid OOXML zip package')
        if ext=='.pdf':
            import fitz
            d=fitz.open(path); r['details']['pages']=d.page_count; r['details']['encrypted']=d.is_encrypted; d.close()
        elif ext=='.pptx':
            with zipfile.ZipFile(path) as z:
                names=z.namelist(); r['details'].update(slides=sum(x.startswith('ppt/slides/slide') and x.endswith('.xml') for x in names),images=sum(x.startswith('ppt/media/') for x in names),embedded=sum(x.startswith('ppt/embeddings/') for x in names))
        elif ext=='.docx':
            from docx import Document
            d=Document(path); r['details'].update(paragraphs=len(d.paragraphs),tables=len(d.tables))
        elif ext=='.xlsx':
            import openpyxl
            w=openpyxl.load_workbook(path,read_only=True,data_only=False); r['details']['sheets']=w.sheetnames; w.close()
        elif ext in {'.png','.jpg','.jpeg','.webp','.gif','.tif','.tiff'}:
            from PIL import Image
            with Image.open(path) as im: r['details'].update(width=im.width,height=im.height,mode=im.mode,frames=getattr(im,'n_frames',1))
        elif ext in {'.md','.txt','.csv','.tsv','.json','.yaml','.yml'}:
            text=path.read_text(encoding='utf-8'); r['details'].update(encoding='utf-8',characters=len(text))
            import re
            if re.search(r'\b\d{3}-\d{2}-\d{4}\b',text): r['sensitive_signals'].append('possible_us_ssn')
            if re.search(r'[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}',text): r['sensitive_signals'].append('email_address_present')
    except Exception as e:
        r['readable']=False; r['errors'].append(f'{type(e).__name__}: {e}')
    return r

def main():
    ap=argparse.ArgumentParser(description=__doc__,formatter_class=argparse.RawDescriptionHelpFormatter); ap.add_argument('inputs',nargs='+'); ap.add_argument('--output','-o',required=True); a=ap.parse_args()
    files=[]
    for raw in a.inputs:
        p=Path(raw)
        files.extend(sorted(x for x in p.rglob('*') if x.is_file())) if p.is_dir() else files.append(p)
    missing=[str(x) for x in files if not x.exists()]
    records=[inspect(x) for x in files if x.exists() and x.is_file()]
    hashes={};
    for x in records:hashes.setdefault(x['sha256'],[]).append(x['path'])
    duplicates=[v for v in hashes.values() if len(v)>1]
    out={'schema_version':'1.1','generated_at':datetime.now(timezone.utc).isoformat(),'files':records,'duplicates':duplicates,'summary':{'count':len(records),'readable':sum(x['readable'] for x in records),'unreadable':sum(not x['readable'] for x in records),'empty':sum(x['empty'] for x in records),'missing':missing}}
    Path(a.output).write_text(json.dumps(out,ensure_ascii=False,indent=2),encoding='utf-8')
    print(json.dumps(out['summary'],ensure_ascii=False)); return 2 if missing or out['summary']['unreadable'] else 0
if __name__=='__main__': raise SystemExit(main())
